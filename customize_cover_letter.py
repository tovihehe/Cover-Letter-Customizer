from docx import Document
from docx2pdf import convert
import shutil
import os

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace text in a paragraph, preserving formatting.
    """
    for run in paragraph.runs:
        for old_word, new_word in replacements.items():
            if old_word in run.text:
                run.text = run.text.replace(old_word, new_word)

def replace_words_in_docx(original_doc_path, replacements, output_doc_path):
    """
    Replace words in a Word document based on a dictionary of replacements.
    
    :param original_doc_path: Path to the input Word document.
    :param replacements: Dictionary of words to be replaced.
    :param output_doc_path: Path to save the modified Word document.
    """

    # Make a copy of the original document to work on
    shutil.copyfile(original_doc_path, output_doc_path)

    # Open the copied Word document
    doc = Document(output_doc_path)

    # Replace words in paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)

    # Replace words in headers and footers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            replace_text_in_paragraph(para, replacements)

        footer = section.footer
        for para in footer.paragraphs:
            replace_text_in_paragraph(para, replacements)

    # Save the modified document
    doc.save(output_doc_path)

    # Convert the modified Word document to PDF
    convert(output_doc_path)
    
    # Delete the modified Word document
    os.remove(output_doc_path)


#### Example usage ####

# Path to the original document
doc_path = 'Template.docx'  

# Path to save the modified document
output_path = 'Cover-Letter.docx' 
          
# Dictionary of words to be replaced          
dictionary_replacement = {
    'DATE': '21st July 2024',
    'POSITION': 'Software Engineer',
    'COMPANY': 'Google',
    'FUNCTION': 'provide innovative software solutions'
}

replace_words_in_docx(doc_path, dictionary_replacement, output_path)
